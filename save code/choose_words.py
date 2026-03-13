import tkinter as tk
from tkinter import scrolledtext, filedialog, messagebox, ttk
import re
import os
from collections import defaultdict
from docx import Document  # 需要安装: pip install python-docx

class EnglishArticleTool:
    def __init__(self, root):
        self.root = root
        self.root.title("英文文章交互工具")
        self.root.geometry("1000x700")
        self.root.minsize(800, 600)
        
        # 预编译正则表达式，提升性能
        # 修改正则表达式，支持连词符连接的单词
        self.word_pattern = re.compile(r"\b[a-zA-Z]+(?:-[a-zA-Z]+)*\b")
        # 用于提取单个单词（包括连词符单词）
        self.single_word_pattern = re.compile(r"[a-zA-Z]+(?:-[a-zA-Z]+)*")
        
        # 简单词汇集合，使用frozenset提升查找性能
        self.simple_words = frozenset({
            "a", "an", "the", "in", "on", "at", "by", "with", "for", "to", "of", "from", 
            "about", "into", "onto", "i", "me", "my", "mine", "you", "your", "yours", 
            "he", "him", "his", "she", "her", "hers", "it", "its", "we", "us", "our", 
            "ours", "they", "them", "their", "theirs", "am", "is", "are", "was", "were", 
            "be", "been", "being", "do", "does", "did", "have", "has", "had", "will", 
            "would", "shall", "should", "can", "could", "may", "might", "must", "this", 
            "that", "these", "those", "here", "there", "now", "then", "and", "or", "but", 
            "so", "not", "no"
        })
        
        self.style = ttk.Style()
        self.style.configure("TButton", font=("微软雅黑", 10))
        self.style.configure("TLabel", font=("微软雅黑", 10))
        
        main_frame = ttk.Frame(root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 按钮区
        button_frame = ttk.Frame(main_frame, padding="5")
        button_frame.pack(fill=tk.X)
        
        ttk.Button(button_frame, text="打开文件", command=self.open_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="保存选中项", command=self.save_selected).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="复制选中项", command=self.copy_selected).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="清除选中", command=self.clear_selected).pack(side=tk.LEFT, padx=5)
        
        # 文本框架
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # 文本框
        self.text_widget = scrolledtext.ScrolledText(
            text_frame, 
            wrap=tk.WORD, 
            font=("Times New Roman", 18),
            padx=10, 
            pady=10
        )
        self.text_widget.pack(fill=tk.BOTH, expand=True)
        
        # 文本选中相关绑定
        self.text_widget.bind("<Button-1>", self.start_selection)
        self.text_widget.bind("<B1-Motion>", self.extend_selection)
        self.text_widget.bind("<ButtonRelease-1>", self.end_selection)
        
        # 变量初始化
        self.current_file = None
        self.word_positions = []
        self.word_counts = defaultdict(int)
        self.selected_items = []
        self.selected_items_set = set()
        self.highlight_tags = {}
        self.selection_in_progress = False
        self.selection_start = None
        self.line_lengths = []
        self.content = ""
        
        # 文本高亮标签配置
        self.text_widget.tag_config("highlight", background="#FFFACD")
        self.text_widget.tag_config("temp_highlight", background="#E6E6FA")
    
    def read_file_content(self, file_path):
        """读取文件内容，支持txt和docx格式"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            elif file_ext == '.docx':
                doc = Document(file_path)
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # 只添加非空段落
                        full_text.append(paragraph.text)
                return '\n'.join(full_text)
            else:
                raise Exception(f"不支持的文件格式: {file_ext}")
        except Exception as e:
            raise Exception(f"读取文件失败: {str(e)}")
    
    def open_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("文本文件", "*.txt"), 
                ("Word文档", "*.docx"),
                ("所有文件", "*.*")
            ],
            title="选择英文文章"
        )
        
        if file_path:
            try:
                self.content = self.read_file_content(file_path)
                
                self.text_widget.delete(1.0, tk.END)
                self.clear_selected()
                
                self.text_widget.insert(tk.END, self.content)
                self.current_file = file_path
                
                self.preprocess_content()
                
                self.root.title(f"英文文章交互工具 - {os.path.basename(file_path)}")
                
            except Exception as e:
                messagebox.showerror("错误", f"无法打开文件: {str(e)}")
    
    def copy_selected(self):
        """复制选中项到剪贴板"""
        if not self.selected_items:
            messagebox.showinfo("提示", "没有选中任何内容")
            return
        
        # 将选中项连接成字符串（每行一项）
        copy_text = '\n'.join(self.selected_items)
        
        # 清除剪贴板并添加新内容
        self.root.clipboard_clear()
        self.root.clipboard_append(copy_text)
        
        # 显示成功消息
        messagebox.showinfo("成功", f"已复制 {len(self.selected_items)} 项到剪贴板")
    
    def save_selected(self):
        if not self.selected_items:
            messagebox.showinfo("提示", "没有选中任何内容")
            return
        
        # 生成默认文件名
        default_filename = "selected_words.txt"
        if self.current_file:
            base_name = os.path.splitext(os.path.basename(self.current_file))[0]
            default_filename = f"{base_name}_selected.txt"
        
        # 选择保存路径
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[
                ("文本文件", "*.txt"), 
                ("Word文档", "*.docx"),
                ("所有文件", "*.*")
            ],
            initialfile=default_filename,
            title="保存选中项"
        )
        
        if file_path:
            try:
                file_ext = os.path.splitext(file_path)[1].lower()
                
                if file_ext == '.txt':
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write('\n'.join(self.selected_items))
                
                elif file_ext == '.docx':
                    doc = Document()
                    doc.add_heading('选中的词汇/短语', level=1)
                    
                    for item in self.selected_items:
                        doc.add_paragraph(item, style='List Bullet')
                    
                    doc.save(file_path)
                
                else:
                    # 默认保存为txt
                    with open(file_path + '.txt', 'w', encoding='utf-8') as file:
                        file.write('\n'.join(self.selected_items))
                    file_path = file_path + '.txt'
                
                messagebox.showinfo("成功", f"已保存 {len(self.selected_items)} 项到文件\n{file_path}")
            
            except Exception as e:
                messagebox.showerror("错误", f"保存文件失败: {str(e)}")
    
    def preprocess_content(self):
        self.word_positions = []
        self.word_counts.clear()
        self.line_lengths = []
        
        # 计算每行长度
        lines = self.content.split('\n')
        for line in lines:
            self.line_lengths.append(len(line) + 1)  # +1 包含换行符
        
        # 提取所有单词（包括带连词符的）
        for match in self.word_pattern.finditer(self.content):
            word = match.group()
            start = match.start()
            end = match.end()
            self.word_positions.append((start, end, word))
            
            # 将带连词符的单词拆分成基本单词来判断是否简单词
            lower_word = word.lower()
            # 如果是带连词符的单词，拆分成基本单词
            if '-' in lower_word:
                base_words = lower_word.split('-')
                # 如果所有基本单词都是简单词，则整个词视为简单词
                if all(w in self.simple_words for w in base_words):
                    continue
            elif lower_word in self.simple_words:
                continue
            
            # 只有包含非简单词的单词才计入词频
            self.word_counts[lower_word] += 1
    
    def start_selection(self, event):
        self.selection_in_progress = True
        self.selection_start = self.get_word_at_position(event)
    
    def extend_selection(self, event):
        if not self.selection_in_progress or self.selection_start is None:
            return
        
        current_word = self.get_word_at_position(event)
        if current_word:
            self.text_widget.tag_remove("temp_highlight", 1.0, tk.END)
            
            start_idx = self.get_text_index(self.selection_start[0])
            end_idx = self.get_text_index(current_word[1])
            self.text_widget.tag_add("temp_highlight", start_idx, end_idx)
    
    def end_selection(self, event):
        if not self.selection_in_progress or self.selection_start is None:
            return
        
        self.selection_in_progress = False
        self.text_widget.tag_remove("temp_highlight", 1.0, tk.END)
        
        current_word = self.get_word_at_position(event)
        if current_word:
            start_pos = min(self.selection_start[0], current_word[0])
            end_pos = max(self.selection_start[1], current_word[1])
            
            start_idx = self.get_text_index(start_pos)
            end_idx = self.get_text_index(end_pos)
            selected_text = self.text_widget.get(start_idx, end_idx).strip()
            
            if not selected_text or selected_text in self.selected_items_set:
                return
            
            # 过滤简单词（包括带连词符的简单词组合）
            selected_words = self.single_word_pattern.findall(selected_text.lower())
            
            # 检查是否为简单词组合
            is_simple = True
            for word in selected_words:
                if '-' in word:
                    # 如果是带连词符的词，检查其组成部分
                    parts = word.split('-')
                    if not all(part in self.simple_words for part in parts):
                        is_simple = False
                        break
                elif word not in self.simple_words:
                    is_simple = False
                    break
            
            if is_simple:
                messagebox.showinfo("提示", f"“{selected_text}”是简单词汇，已自动排除")
                return
            
            self.selected_items.append(selected_text)
            self.selected_items_set.add(selected_text)
            
            self.highlight_selected_text(selected_text)
    
    def get_word_at_position(self, event):
        try:
            index = self.text_widget.index(f"@{event.x},{event.y}")
            line, column = map(int, index.split('.'))
            
            # 计算字符位置
            char_pos = sum(self.line_lengths[:line-1]) + column
            
            # 二分查找
            left, right = 0, len(self.word_positions) - 1
            while left <= right:
                mid = (left + right) // 2
                start, end, word = self.word_positions[mid]
                
                if start <= char_pos < end:
                    return (start, end, word)
                elif char_pos < start:
                    right = mid - 1
                else:
                    left = mid + 1
            
            return None
        except:
            return None
    
    def get_text_index(self, char_pos):
        left, right = 0, len(self.line_lengths) - 1
        while left <= right:
            mid = (left + right) // 2
            line_total = sum(self.line_lengths[:mid+1])
            
            if line_total > char_pos:
                right = mid - 1
            else:
                left = mid + 1
        
        line = left + 1
        prev_total = sum(self.line_lengths[:left]) if left > 0 else 0
        column = char_pos - prev_total
        
        return f"{line}.{column}"
    
    def calculate_occurrences(self, text):
        """计算选中文本在文章中的出现次数"""
        if self.single_word_pattern.fullmatch(text) and '-' not in text:
            # 单个单词（不带连词符）：直接用预缓存的词频
            lower_text = text.lower()
            return self.word_counts.get(lower_text, 0)
        
        # 短语或带连词符的单词：遍历文章查找完整匹配
        count = 0
        start = 0
        text_len = len(text)
        content = self.content
        
        while True:
            pos = content.find(text, start)
            if pos == -1:
                break
            
            # 检查是否为完整匹配（前后非字母数字和连词符）
            prev_char = content[pos-1] if pos > 0 else ' '
            next_char = content[pos+text_len] if pos+text_len < len(content) else ' '
            
            if (not prev_char.isalnum() and prev_char != '-' or pos == 0) and \
               (not next_char.isalnum() and next_char != '-' or pos+text_len == len(content)):
                # 过滤仅包含简单词的短语
                match_words = self.single_word_pattern.findall(content[pos:pos+text_len].lower())
                
                is_simple = True
                for word in match_words:
                    if '-' in word:
                        parts = word.split('-')
                        if not all(part in self.simple_words for part in parts):
                            is_simple = False
                            break
                    elif word not in self.simple_words:
                        is_simple = False
                        break
                
                if not is_simple:
                    count += 1
            
            start = pos + 1
        
        return count
    
    def highlight_selected_text(self, text):
        tag = f"highlight_{len(self.highlight_tags)}"
        self.highlight_tags[text] = tag
        self.text_widget.tag_config(tag, background="#FFFACD")
        
        content = self.content
        start = 0
        text_len = len(text)
        
        while True:
            pos = content.find(text, start)
            if pos == -1:
                break
            
            # 检查完整匹配（考虑连词符）
            prev_char = content[pos-1] if pos > 0 else ' '
            next_char = content[pos+text_len] if pos+text_len < len(content) else ' '
            
            if (not prev_char.isalnum() and prev_char != '-' or pos == 0) and \
               (not next_char.isalnum() and next_char != '-' or pos+text_len == len(content)):
                # 检查是否包含非简单词
                match_words = self.single_word_pattern.findall(content[pos:pos+text_len].lower())
                
                is_simple = True
                for word in match_words:
                    if '-' in word:
                        parts = word.split('-')
                        if not all(part in self.simple_words for part in parts):
                            is_simple = False
                            break
                    elif word not in self.simple_words:
                        is_simple = False
                        break
                
                if not is_simple:
                    start_idx = self.get_text_index(pos)
                    end_idx = self.get_text_index(pos + text_len)
                    self.text_widget.tag_add(tag, start_idx, end_idx)
            
            start = pos + 1
    
    def clear_selected(self):
        # 移除所有永久高亮标签
        for tag in self.highlight_tags.values():
            self.text_widget.tag_remove(tag, 1.0, tk.END)
        
        # 重置选中项相关变量
        self.selected_items = []
        self.selected_items_set = set()
        self.highlight_tags = {}

if __name__ == "__main__":
    root = tk.Tk()
    app = EnglishArticleTool(root)
    root.mainloop()